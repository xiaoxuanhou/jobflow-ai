"""
Utility functions for jobflow-ai.

Handles DOCX rendering, PDF conversion, file output, and debugging.
"""

import json
import logging
import re
from datetime import datetime, date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


def slugify(text: str) -> str:
    """
    Convert text to a safe folder/file name.

    Args:
        text: Text to slugify (e.g., "Applied Scientist - Anthropic")

    Returns:
        Safe slug (e.g., "applied-scientist-anthropic")
    """
    # Convert to lowercase
    text = text.lower()
    # Replace spaces and underscores with hyphens
    text = re.sub(r'[\s_]+', '-', text)
    # Remove non-alphanumeric characters (except hyphens)
    text = re.sub(r'[^a-z0-9\-]', '', text)
    # Remove multiple consecutive hyphens
    text = re.sub(r'-+', '-', text)
    # Remove leading/trailing hyphens
    text = text.strip('-')
    # Limit length
    return text[:60]


# =============================================================================
# DOCX Rendering Functions
# =============================================================================


def _add_horizontal_rule(paragraph):
    """Add a horizontal rule (bottom border) below a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')  # 1/8 pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, text: str, url: str):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '004F90')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _format_text_with_bold(paragraph, text: str, base_bold: bool = False):
    """
    Format text with **bold** markdown converted to actual bold.

    Args:
        paragraph: The paragraph to add text to
        text: Text that may contain **bold** markers
        base_bold: Whether the base text should be bold
    """
    # Split on **bold** markers
    parts = re.split(r'\*\*([^*]+)\*\*', text)

    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        # Odd indices are the bold parts (between **)
        if i % 2 == 1 or base_bold:
            run.bold = True


def render_docx(tailored_content: dict, output_path: str) -> str:
    """
    Generate a DOCX resume from tailored content.

    Args:
        tailored_content: Dict with contact, summary, experiences, education, skills, publications
        output_path: Path to save the .docx file

    Returns:
        Path to generated .docx file
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    contact = tailored_content.get('contact', {})
    summary = tailored_content.get('summary', '')
    experiences = tailored_content.get('experiences', [])
    education = tailored_content.get('education', [])
    skills = tailored_content.get('skills', [])
    publications = tailored_content.get('publications', [])

    # =========================================================================
    # HEADER - Name (centered, large, bold)
    # =========================================================================
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(contact.get('name', ''))
    name_run.bold = True
    name_run.font.size = Pt(24)

    # Contact info on one line with separators
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_items = []
    if contact.get('email'):
        contact_items.append(('email', contact['email']))
    if contact.get('phone'):
        contact_items.append(('text', contact['phone']))
    if contact.get('linkedin'):
        contact_items.append(('link', 'LinkedIn', contact['linkedin']))
    if contact.get('github'):
        contact_items.append(('link', 'GitHub', contact['github']))
    if contact.get('work_authorization'):
        contact_items.append(('text', contact['work_authorization']))

    for i, item in enumerate(contact_items):
        if i > 0:
            contact_para.add_run(' | ')

        if item[0] == 'email':
            _add_hyperlink(contact_para, item[1], f"mailto:{item[1]}")
        elif item[0] == 'link':
            _add_hyperlink(contact_para, item[1], item[2])
        else:
            contact_para.add_run(item[1])

    # =========================================================================
    # SUMMARY - Centered, italicized
    # =========================================================================
    if summary:
        summary_para = doc.add_paragraph()
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary_run = summary_para.add_run(summary)
        summary_run.italic = True

    # =========================================================================
    # PROFESSIONAL EXPERIENCE
    # =========================================================================
    if experiences:
        section_para = doc.add_paragraph()
        section_run = section_para.add_run('PROFESSIONAL EXPERIENCE')
        section_run.bold = True
        section_run.font.size = Pt(12)
        _add_horizontal_rule(section_para)

        for exp in experiences:
            # Role, Company (right-aligned dates) using a table
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            table.allow_autofit = True

            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            # Left: Role, Company
            left_para = left_cell.paragraphs[0]
            role_run = left_para.add_run(exp.get('role', ''))
            role_run.bold = True
            left_para.add_run(f", {exp.get('company', '')}")

            # Right: Dates
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_str = exp.get('start_date', '')
            if exp.get('end_date') and exp.get('end_date') != exp.get('start_date'):
                date_str += f" – {exp.get('end_date', '')}"
            right_para.add_run(date_str)

            # Location (italic)
            if exp.get('location'):
                loc_para = doc.add_paragraph()
                loc_run = loc_para.add_run(exp.get('location', ''))
                loc_run.italic = True

            # Bullet points
            for bullet in exp.get('bullets', []):
                bullet_para = doc.add_paragraph(style='List Bullet')
                _format_text_with_bold(bullet_para, bullet)

    # =========================================================================
    # EDUCATION
    # =========================================================================
    if education:
        section_para = doc.add_paragraph()
        section_run = section_para.add_run('EDUCATION')
        section_run.bold = True
        section_run.font.size = Pt(12)
        _add_horizontal_rule(section_para)

        for edu in education:
            # Institution (right-aligned location)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True

            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            left_para = left_cell.paragraphs[0]
            inst_run = left_para.add_run(edu.get('institution', ''))
            inst_run.bold = True

            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para.add_run(edu.get('location', ''))

            # Degree (right-aligned graduation)
            table2 = doc.add_table(rows=1, cols=2)
            table2.autofit = True

            left_cell2 = table2.rows[0].cells[0]
            right_cell2 = table2.rows[0].cells[1]

            left_para2 = left_cell2.paragraphs[0]
            degree_run = left_para2.add_run(edu.get('degree', ''))
            degree_run.italic = True
            if edu.get('focus'):
                left_para2.add_run(f" (focus: {edu.get('focus')})")

            right_para2 = right_cell2.paragraphs[0]
            right_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para2.add_run(edu.get('graduation', ''))

            # Advisors
            if edu.get('advisors'):
                advisors_para = doc.add_paragraph()
                advisors_para.add_run('Advisors: ')
                advisors_para.add_run(', '.join(edu.get('advisors', [])))

    # =========================================================================
    # SELECTED PUBLICATIONS
    # =========================================================================
    if publications:
        section_para = doc.add_paragraph()
        section_run = section_para.add_run('SELECTED PUBLICATIONS')
        section_run.bold = True
        section_run.font.size = Pt(12)
        _add_horizontal_rule(section_para)

        for pub in publications:
            pub_para = doc.add_paragraph(style='List Bullet')

            # First author indicator
            if pub.get('first_author'):
                fa_run = pub_para.add_run('First-author ')
                fa_run.bold = True

            # Status and venue
            status_text = f"({pub.get('status', '')}"
            if pub.get('venue'):
                status_text += f", "
                pub_para.add_run(status_text)
                venue_run = pub_para.add_run(pub.get('venue'))
                venue_run.italic = True
                pub_para.add_run('): ')
            else:
                pub_para.add_run(status_text + '): ')

            # Title
            pub_para.add_run(pub.get('short_title', ''))

            # Link
            if pub.get('url'):
                pub_para.add_run(' [')
                _add_hyperlink(pub_para, pub.get('url_label', 'Link'), pub.get('url'))
                pub_para.add_run(']')

    # =========================================================================
    # TECHNICAL SKILLS
    # =========================================================================
    if skills:
        section_para = doc.add_paragraph()
        section_run = section_para.add_run('TECHNICAL SKILLS')
        section_run.bold = True
        section_run.font.size = Pt(12)
        _add_horizontal_rule(section_para)

        for category in skills:
            skill_para = doc.add_paragraph()
            cat_run = skill_para.add_run(f"{category.get('name', '')}:")
            cat_run.bold = True
            skill_para.add_run(' ')
            skill_list = category.get('skill_list', [])
            if isinstance(skill_list, list):
                skill_para.add_run(', '.join(skill_list))
            else:
                skill_para.add_run(str(skill_list))

    # Save document
    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path: str) -> Optional[str]:
    """
    Convert DOCX to PDF using docx2pdf.

    Uses Word on Windows, LibreOffice on Linux/Mac.

    Args:
        docx_path: Path to the .docx file

    Returns:
        Path to PDF if successful, None otherwise
    """
    try:
        from docx2pdf import convert
        pdf_path = docx_path.replace('.docx', '.pdf')
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        logger.warning("docx2pdf not installed. Install with: pip install docx2pdf")
        return None
    except Exception as e:
        logger.error(f"PDF conversion failed: {e}")
        logger.info("On Windows, ensure Microsoft Word is installed.")
        logger.info("On Linux/Mac, ensure LibreOffice is installed.")
        return None


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract all text content from a DOCX file.

    Args:
        docx_path: Path to the .docx file

    Returns:
        Extracted text content
    """
    doc = Document(docx_path)
    text_parts = []

    for para in doc.paragraphs:
        text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    return '\n'.join(text_parts)


def render_outreach_md(targets: list, scripts: dict, company: str, role: str) -> str:
    """
    Format outreach plan as markdown with copy-paste scripts.

    Args:
        targets: List of outreach target dicts
        scripts: Dict mapping names to their scripts
        company: Company name
        role: Role title

    Returns:
        Formatted markdown string
    """
    lines = [
        f"# Outreach Plan: {role} @ {company}",
        "",
    ]

    for i, target in enumerate(targets, 1):
        name = target.get('name', 'Unknown')
        lines.extend([
            f"## Contact {i}: {name} ({target.get('connection_type', 'Unknown')} - Priority {target.get('priority', i)})",
            f"**LinkedIn**: {target.get('linkedin_url', 'N/A')}",
            f"**Title**: {target.get('title', 'N/A')}",
            f"**Connection Angle**: {target.get('connection_angle', 'N/A')}",
            "",
        ])

        if name in scripts:
            script = scripts[name]
            if 'connection_request' in script:
                lines.extend([
                    "### Connection Request (copy/paste):",
                    f"> {script['connection_request']}",
                    "",
                ])
            if 'follow_up_message' in script:
                lines.extend([
                    "### Follow-up Message (after accepted):",
                    f"> {script['follow_up_message']}",
                    "",
                ])

        lines.append("---")
        lines.append("")

    # Add outreach sequence if present
    if 'outreach_sequence' in scripts:
        lines.extend([
            "## Recommended Outreach Sequence",
        ])
        for step in scripts['outreach_sequence']:
            lines.append(f"- {step}")
        lines.append("")

    return '\n'.join(lines)


def save_model_inputs(
    job_id: str,
    prompt: str,
    response: dict,
    output_dir: str,
    input_tokens: int = 0,
    output_tokens: int = 0,
    cost: float = 0.0,
    model: str = "claude-sonnet-4-20250514"
) -> None:
    """
    Save model inputs alongside outputs for debugging/auditing.

    Creates:
    - {output_dir}/debug/prompt.md      - Full prompt sent to model
    - {output_dir}/debug/response.json  - Raw model response
    - {output_dir}/debug/metadata.json  - Timestamps, token counts, costs

    Args:
        job_id: Job identifier
        prompt: Full prompt sent to model
        response: Parsed response dict
        output_dir: Base output directory for this job
        input_tokens: Number of input tokens used
        output_tokens: Number of output tokens used
        cost: Estimated cost in USD
        model: Model identifier
    """
    debug_dir = Path(output_dir) / "debug"
    debug_dir.mkdir(parents=True, exist_ok=True)

    # Save prompt
    (debug_dir / "prompt.md").write_text(prompt, encoding='utf-8')

    # Save response
    (debug_dir / "response.json").write_text(
        json.dumps(response, indent=2, ensure_ascii=False),
        encoding='utf-8'
    )

    # Save metadata
    metadata = {
        "job_id": job_id,
        "timestamp": datetime.now().isoformat(),
        "model": model,
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "cost_usd": round(cost, 4),
    }
    (debug_dir / "metadata.json").write_text(
        json.dumps(metadata, indent=2),
        encoding='utf-8'
    )


def generate_daily_summary(results: list, output_dir: str) -> str:
    """
    Create daily digest of all processed jobs.

    Args:
        results: List of ProcessedOutput dicts
        output_dir: Base output directory

    Returns:
        Path to the generated summary file
    """
    today = date.today().isoformat()
    lines = [
        f"# jobflow-ai Results - {today}",
        "",
        f"## Top {len(results)} Ranked Opportunities",
        "",
    ]

    for i, result in enumerate(results, 1):
        analysis = result.get('analysis', {})
        company = analysis.get('company', 'Unknown')
        role = analysis.get('role', 'Unknown')
        job_url = result.get('job_url', '#')

        lines.extend([
            f"### #{i}: {role} - {company}",
            f"**Link**: {job_url}",
            "",
        ])

        # Key requirements
        if analysis.get('key_requirements'):
            lines.append("**Key Requirements**: " + ", ".join(analysis['key_requirements'][:3]))

        # Candidate strengths
        if analysis.get('candidate_strengths'):
            lines.append("**Your Strengths**: " + ", ".join(analysis['candidate_strengths'][:3]))

        lines.append("")

        # Files generated
        job_slug = result.get('slug', slugify(f"{company}-{role}"))
        lines.extend([
            f"**Files**: `{job_slug}/resume.docx`, `{job_slug}/outreach.md`",
            "",
        ])

        # Top contacts
        targets = result.get('outreach_targets', [])
        if targets:
            lines.append("**Top Contacts**:")
            lines.append("| Priority | Name | Type | LinkedIn |")
            lines.append("|----------|------|------|----------|")
            for target in targets[:3]:
                lines.append(
                    f"| {target.get('priority', '-')} | {target.get('name', 'N/A')} | "
                    f"{target.get('connection_type', 'N/A')} | "
                    f"[Profile]({target.get('linkedin_url', '#')}) |"
                )
            lines.append("")

        lines.append("---")
        lines.append("")

    # Save summary
    summary_path = Path(output_dir) / today / "daily_summary.md"
    summary_path.parent.mkdir(parents=True, exist_ok=True)
    summary_path.write_text('\n'.join(lines), encoding='utf-8')

    return str(summary_path)


def load_yaml(path: str) -> dict:
    """Load a YAML file."""
    import yaml
    with open(path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f) or {}


def save_json(data: dict, path: str) -> None:
    """Save data as JSON."""
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    Path(path).write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')


def load_json(path: str) -> dict:
    """Load a JSON file."""
    return json.loads(Path(path).read_text(encoding='utf-8'))


def get_output_dir(company: str, role: str, base_dir: str = "outputs") -> Path:
    """
    Get the output directory for a job.

    Args:
        company: Company name
        role: Role title
        base_dir: Base output directory

    Returns:
        Path to the job-specific output directory
    """
    today = date.today().isoformat()
    slug = slugify(f"{company}-{role}")
    path = Path(base_dir) / today / slug
    path.mkdir(parents=True, exist_ok=True)
    return path
